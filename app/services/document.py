"""Multi-format resume document text extractor (PDF, DOCX, TXT).

Includes magic-byte format validation and robust fallback extraction engines.
"""

from __future__ import annotations

import io
import logging
import zipfile
import xml.etree.ElementTree as ET

from pypdf import PdfReader
import docx

from app.config import settings
from app.core.security import InputSanitizer

logger = logging.getLogger("skillbridge.document")

PDF_MAGIC = b"%PDF-"
DOCX_MAGIC = b"PK\x03\x04"


def validate_document_bytes(raw_bytes: bytes, filename: str = "") -> tuple[bool, str, str]:
    """Validate file format, magic bytes, and size bounds.
    
    Returns (is_valid, detected_type, error_message).
    """
    if not raw_bytes:
        return False, "unknown", "Uploaded file is empty."

    if len(raw_bytes) > settings.max_upload_bytes:
        max_mb = settings.max_upload_bytes // (1024 * 1024)
        return False, "unknown", f"File exceeds maximum allowed size of {max_mb} MB."

    fn_lower = filename.lower().strip()
    header_1024 = raw_bytes[:1024]

    # 1. PDF Detection
    if PDF_MAGIC in header_1024 or fn_lower.endswith(".pdf"):
        if PDF_MAGIC not in header_1024:
            return False, "pdf", "Invalid PDF file: Missing standard %PDF- magic header signature."
        return True, "pdf", "Valid PDF"

    # 2. DOCX Detection
    if header_1024.startswith(DOCX_MAGIC) or fn_lower.endswith(".docx"):
        if not header_1024.startswith(DOCX_MAGIC):
            return False, "docx", "Invalid DOCX file: Missing standard ZIP/DOCX header signature."
        # Verify it contains word/document.xml
        try:
            with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
                if "word/document.xml" not in zf.namelist():
                    return False, "docx", "Invalid DOCX file: Missing word/document.xml content."
        except Exception:
            return False, "docx", "Corrupted or invalid DOCX archive."
        return True, "docx", "Valid DOCX"

    # 3. Plain Text Detection
    if fn_lower.endswith(".txt") or fn_lower.endswith(".md"):
        try:
            raw_bytes.decode("utf-8")
            return True, "text", "Valid Text"
        except UnicodeDecodeError:
            try:
                raw_bytes.decode("latin-1")
                return True, "text", "Valid Text"
            except Exception:
                return False, "text", "Unsupported text encoding."

    return False, "unknown", "Unsupported file type. Please upload a PDF (.pdf) or Word document (.docx)."


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF document with page bounds."""
    pages: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes), strict=False)
        for page in reader.pages[: settings.max_pdf_pages]:
            try:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())
            except Exception:
                continue
    except Exception as exc:
        logger.warning("PDF extraction warning: %s", exc)

    extracted = "\n\n".join(pages)
    return InputSanitizer.sanitize_text(extracted, max_chars=settings.max_text_chars)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX document with python-docx and XML fallback."""
    paragraphs: list[str] = []

    # Method A: python-docx library
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
    except Exception as exc:
        logger.debug("python-docx reader fallback triggered: %s", exc)

    # Method B: Direct XML parsing fallback if Method A yielded empty
    if not paragraphs:
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                xml_content = zf.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                # Word XML namespace
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for node in tree.iterfind(".//w:p", namespaces):
                    texts = [t.text for t in node.iterfind(".//w:t", namespaces) if t.text]
                    line = "".join(texts).strip()
                    if line:
                        paragraphs.append(line)
        except Exception as exc:
            logger.warning("XML fallback DOCX extraction warning: %s", exc)

    extracted = "\n\n".join(paragraphs)
    return InputSanitizer.sanitize_text(extracted, max_chars=settings.max_text_chars)


def extract_document_text(file_bytes: bytes, filename: str = "") -> tuple[str, str]:
    """Validate and extract text from an uploaded document (PDF, DOCX, TXT).
    
    Returns (extracted_text, file_type).
    """
    is_valid, file_type, err_msg = validate_document_bytes(file_bytes, filename=filename)
    if not is_valid:
        raise ValueError(err_msg)

    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        text = extract_text_from_docx(file_bytes)
    elif file_type == "text":
        try:
            text = file_bytes.decode("utf-8")
        except Exception:
            text = file_bytes.decode("latin-1", errors="ignore")
        text = InputSanitizer.sanitize_text(text, max_chars=settings.max_text_chars)
    else:
        text = ""

    if not text:
        text = "(No extractable text found in the uploaded document. The file may contain scanned images or empty pages.)"

    return text, file_type
